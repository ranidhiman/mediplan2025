"""
rag_api.py — Tiered document extraction pipeline
Tier 0: Regex  |  Tier 1: spaCy NLP  |  Tier 2: Ollama (adaptive schema)
Storage: SQLite (schemas + chunks) + ChromaDB (embeddings for RAG)
"""
import json, re, uuid, sqlite3
from pathlib import Path
from datetime import datetime

from flask import Flask, request, jsonify
from flask_cors import CORS
import pdfplumber
import spacy
import chromadb
import requests
import docx
import openpyxl

app = Flask(__name__)
CORS(app)

# ── spaCy ──────────────────────────────────────────────────────────────────────
print("Loading spaCy...")
nlp = spacy.load("en_core_web_sm")

# ── ChromaDB ───────────────────────────────────────────────────────────────────
print("Initializing ChromaDB...")
chroma_client = chromadb.PersistentClient(path="./chroma_db")

# Use the same embedding model already on disk
from chromadb.utils import embedding_functions
_ef = embedding_functions.SentenceTransformerEmbeddingFunction(
    model_name="./models/embed/all-MiniLM-L6-v2"
)
chroma_col = chroma_client.get_or_create_collection("documents", embedding_function=_ef)

# ── SQLite ─────────────────────────────────────────────────────────────────────
DB_PATH = "./extraction.db"

def _db():
    return sqlite3.connect(DB_PATH)

def _init_db():
    with _db() as con:
        con.execute("""CREATE TABLE IF NOT EXISTS documents (
            id          TEXT PRIMARY KEY,
            file_name   TEXT,
            file_path   TEXT,
            doc_type    TEXT,
            schema_json TEXT,
            extracted_at TEXT
        )""")
        con.execute("""CREATE TABLE IF NOT EXISTS chunks (
            id          TEXT PRIMARY KEY,
            document_id TEXT,
            page_num    INTEGER,
            text        TEXT
        )""")
        con.commit()

_init_db()

# ── Ollama client ──────────────────────────────────────────────────────────────
OLLAMA_URL   = "http://localhost:11434"
OLLAMA_MODEL = "llama3.1:8b"

def ollama_chat(prompt, system="Output valid JSON only. No markdown. No explanation.", max_tokens=2048):
    full_prompt = f"{system}\n\n{prompt}"
    try:
        print(f"[Ollama] sending request ({len(full_prompt)} chars)...")
        resp = requests.post(f"{OLLAMA_URL}/api/generate", json={
            "model": OLLAMA_MODEL,
            "prompt": full_prompt,
            "stream": False,
            "options": {"temperature": 0.0, "num_predict": max_tokens},
        }, timeout=300)
        resp.raise_for_status()
        content = resp.json().get("response", "{}")
        print(f"[Ollama] response ({len(content)} chars): {content[:300]}")
        return content
    except Exception as e:
        print(f"[Ollama] error: {e}")
        return "{}"

# ── Document readers ───────────────────────────────────────────────────────────
def read_pdf(path):
    pages = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({"page_num": i + 1, "text": text})
    return pages

def read_docx(path):
    doc = docx.Document(path)
    sections, current, num = [], [], 1
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and current:
            sections.append({"page_num": num, "text": "\n".join(current)})
            num += 1
            current = [para.text]
        elif para.text.strip():
            current.append(para.text)
    if current:
        sections.append({"page_num": num, "text": "\n".join(current)})
    if not sections:
        all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        sections = [{"page_num": 1, "text": all_text}]
    return sections

def read_xlsx(path):
    """Each sheet is treated as a separate page."""
    wb = openpyxl.load_workbook(path, data_only=True)
    pages = []
    for i, sheet in enumerate(wb.worksheets):
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if any(c is not None for c in row):
                rows.append("\t".join("" if c is None else str(c) for c in row))
        if rows:
            pages.append({"page_num": i + 1, "text": f"[Sheet: {sheet.title}]\n" + "\n".join(rows)})
    return pages

def read_txt(path):
    text = Path(path).read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    pages = []
    for i in range(0, max(len(lines), 1), 100):
        chunk = "\n".join(lines[i:i + 100]).strip()
        if chunk:
            pages.append({"page_num": i // 100 + 1, "text": chunk})
    return pages or [{"page_num": 1, "text": text}]

def read_document(file_path):
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":                       return read_pdf(file_path)
    if ext in (".docx", ".doc"):            return read_docx(file_path)
    if ext in (".xlsx", ".xls"):            return read_xlsx(file_path)
    if ext in (".txt", ".md", ".log", ".csv"): return read_txt(file_path)
    return []

# ── Tier 0: Regex ──────────────────────────────────────────────────────────────
SKIP_TITLE = {
    "unclassified", "secret", "top secret", "confidential", "fouo", "cui",
    "noforn", "restricted", "for official use only", "unclassified//fouo",
}
SKIP_TITLE_RE = re.compile(
    r"^(copy\s+\d+\s+of\s+\d+|unclassified|secret|top secret|confidential|"
    r"fouo|cui|noforn|restricted|for official use only)$",
    re.IGNORECASE,
)
MONTH_MAP = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
    "january": 1, "february": 2, "march": 3, "april": 4, "june": 6,
    "july": 7, "august": 8, "september": 9, "october": 10, "november": 11, "december": 12,
}
MON     = r"Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec"
MONFULL = r"January|February|March|April|May|June|July|August|September|October|November|December"

def _parse_date(raw):
    from datetime import date as _d
    raw = raw.strip()
    for pat in [
        rf"^(\d{{1,2}})({MON})(\d{{2,4}})$",
        rf"^(\d{{1,2}})\s+({MON}|{MONFULL})\s+(\d{{4}})$",
    ]:
        m = re.match(pat, raw, re.IGNORECASE)
        if m:
            day = int(m.group(1))
            mon_s = MONTH_MAP.get(m.group(2).lower()[:3])
            yr  = int(m.group(3))
            if mon_s is None: continue
            if yr < 100: yr += 2000
            try:    return _d(yr, mon_s, day).strftime("%Y-%m-%dT00:00:00.000Z")
            except: continue
    return None

def tier0_regex(text):
    fields = {}

    # Date — try three patterns in priority order
    date_found = False
    for pat, fmt in [
        (rf"\b(\d{{1,2}}(?:{MON})\d{{2,4}})\b",               "short"),
        (rf"\b(\d{{6}}Z?\s*(?:{MON})\s*\d{{2,4}})\b",         "dtg"),
        (rf"\b(\d{{1,2}}\s+(?:{MON}|{MONFULL})\s+\d{{4}})\b", "long"),
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            raw = m.group(1).strip()
            if fmt == "dtg":
                raw = re.sub(r"^(\d{2})\d{4}Z?\s*", r"\1", raw).strip()
            parsed = _parse_date(raw)
            fields["date"] = {
                "value": parsed or m.group(1).strip(),
                "confidence": 1.0 if parsed else 0.6,
                "method": "regex",
            }
            date_found = True
            break
    if not date_found:
        fields["date"] = {"value": None, "confidence": 0.0, "method": "regex"}

    # MGRS grid
    ms = re.search(r"\b(\d{2}[A-Z]{2,3}\s+\d{4,5}\s+\d{4,5})\b", text)
    mc = re.search(r"\b(\d{2}[A-Z]{2,3}\d{8,10})\b", text)
    if ms:   fields["mgrsGrid"] = {"value": ms.group(1), "confidence": 1.0, "method": "regex"}
    elif mc: fields["mgrsGrid"] = {"value": mc.group(1), "confidence": 0.9, "method": "regex"}
    else:    fields["mgrsGrid"] = {"value": None, "confidence": 0.0, "method": "regex"}

    # Casualty
    cm = re.search(r"\bCAx(\d+)\b", text, re.IGNORECASE) or \
         re.search(r"\b(\d+)\s+(?:WIA|KIA|casualties|wounded|killed)\b", text, re.IGNORECASE)
    if cm:
        is_cax = bool(re.search(r"CAx", cm.group(0), re.IGNORECASE))
        fields["casualty"] = {
            "value": f"CAx{cm.group(1)}" if is_cax else cm.group(0).strip(),
            "confidence": 1.0 if is_cax else 0.75,
            "method": "regex",
        }
    else:
        fields["casualty"] = {"value": None, "confidence": 0.0, "method": "regex"}

    # Title — ranked priority (highest confidence wins)
    hm = re.search(r"^#{1,3}\s+(.+)$", text, re.MULTILINE)
    if hm:
        fields["title"] = {"value": hm.group(1).strip(), "confidence": 0.95, "method": "regex"}
    else:
        # Primary doc-type headings (AAR, SITREP, WARNO, etc.) at start of a line —
        # checked BEFORE the ANNEX/APPENDIX pattern so a stray "OPORD 07-21" reference
        # inside an AAR doesn't override the actual document title.
        primary_heading = re.search(
            r"^((?:AFTER\s+ACTION\s+(?:REPORT|REVIEW)|SITUATION\s+REPORT|WARNING\s+ORDER|"
            r"FRAGMENTARY\s+ORDER|OPERATION\s+PLAN|OPERATIONS\s+ORDER)"
            r"(?:\s*[\(\-–]\s*[^\n]{0,80})?)",
            text, re.IGNORECASE | re.MULTILINE,
        )
        if primary_heading:
            fields["title"] = {"value": re.sub(r"\s+", " ", primary_heading.group(1)).strip(), "confidence": 0.93, "method": "regex"}
        else:
            # ANNEX/APPENDIX cross-reference line
            bm = re.search(
                r"((?:APPENDIX|ANNEX)\s+\w+\s*\([^)]+\)\s+TO\s+"
                r"(?:ANNEX\s+\w+\s*\([^)]+\)\s*OF\s+)?(?:OPERATION\s+ORDER|OPORD|FRAGO|WARNO|OPLAN)"
                r"(?:\s+[\w\-]+(?:\s+\([^)]+\))?(?:\s*-[^\n]+)?)?)",
                text, re.IGNORECASE | re.DOTALL,
            )
            if bm:
                fields["title"] = {"value": re.sub(r"\s+", " ", bm.group(1)).strip(), "confidence": 0.88, "method": "regex"}
            else:
                dm = re.search(
                    r"^((?:OPORD|AAR|FRAGO|WARNO|OPLAN)\s+[\w\-\/]+(?:\s+[\w\-\/]+)?)",
                    text, re.IGNORECASE | re.MULTILINE,
                )
                if dm:
                    fields["title"] = {"value": dm.group(1).strip(), "confidence": 0.85, "method": "regex"}
                else:
                    fl = next((
                        l.strip() for l in text.splitlines()
                        if l.strip() and len(l.strip()) > 4
                        and l.strip().lower() not in SKIP_TITLE
                        and not SKIP_TITLE_RE.match(l.strip())
                    ), None)
                    fields["title"] = {"value": fl, "confidence": 0.5 if fl else 0.0, "method": "regex-fallback"}

    return fields

# ── Tier 1: spaCy ──────────────────────────────────────────────────────────────
TIER1_THRESHOLD = 0.7

def _clean_ents(ent_list, min_len=3, max_len=60):
    """Remove noise: short tokens, all-caps acronyms that aren't real names, (U) prefixes."""
    out = []
    seen = set()
    for e in ent_list:
        e = re.sub(r'^\(U\)\s*', '', e).strip()
        if len(e) < min_len or len(e) > max_len:
            continue
        if e.lower() in seen:
            continue
        # Skip lines that look like sentence fragments or section headers
        if re.search(r'[:\/\n]', e):
            continue
        seen.add(e.lower())
        out.append(e)
    return out

def tier1_spacy(text, fields):
    doc = nlp(text[:100_000])

    ents = {}
    for label in ("DATE", "PERSON", "ORG", "GPE", "LOC", "FAC", "CARDINAL", "EVENT"):
        raw = list(dict.fromkeys(e.text for e in doc.ents if e.label_ == label))
        ents[label] = _clean_ents(raw)

    updated = dict(fields)

    # Promote date if regex missed it
    if updated.get("date", {}).get("confidence", 0) < TIER1_THRESHOLD and ents["DATE"]:
        updated["date"] = {"value": ents["DATE"][0], "confidence": 0.6, "method": "spacy"}

    # Roles/section headers that spaCy mis-tags as PERSON
    NON_PERSON_FRAGMENTS = {
        'surgeon', 'officer', 'commander', 'commander,', 'annex', 'lessons',
        'appendix', 'sustainment', 'support', 'authenticator', 'prepared',
        'contact', 'team', 'company', 'battalion', 'brigade', 'division',
    }
    def _looks_like_person(name):
        if " " not in name:
            return False
        words = name.lower().split()
        # Reject if any word is a known non-person fragment
        if any(w.rstrip('.,') in NON_PERSON_FRAGMENTS for w in words):
            return False
        # Must have at least one word that looks like a proper first/last name
        # (capitalised, alphabetic, 2+ chars)
        return any(w[0].isupper() and w.isalpha() and len(w) > 1 for w in name.split())

    real_people = [p for p in ents["PERSON"] if _looks_like_person(p)]
    if real_people:
        updated["personnel"] = {"value": real_people, "confidence": 0.75, "method": "spacy"}

    # ORG: keep only entries that look like real unit/org names
    real_orgs = [o for o in ents["ORG"] if len(o) > 5 and not o.startswith("U)")]
    if real_orgs:
        updated["units"] = {"value": real_orgs[:20], "confidence": 0.75, "method": "spacy"}

    # LOC/FAC: exclude things that look like roles or section headers
    NON_LOC_FRAGMENTS = {'surgeon', 'officer', 'commander', 'lessons', 'appendix', 'annex', 'sustainment'}
    locs = [l for l in ents["GPE"] + ents["LOC"] + ents["FAC"]
            if not any(f in l.lower() for f in NON_LOC_FRAGMENTS)]
    if locs:
        updated["locations"] = {"value": list(dict.fromkeys(locs))[:15], "confidence": 0.75, "method": "spacy"}

    return updated, ents

# ── Tier 2: Ollama — format-aware adaptive schema ─────────────────────────────
TIER2_THRESHOLD = 0.7
HALLUCINATED = {"38T MN 12345 67890", "38TMN1234567890", "CAx3", "CAx0", "OPORD 1-26 IRON SHIELD"}

# ── Document type detection ────────────────────────────────────────────────────
DOC_TYPE_PATTERNS = [
    # More-specific / primary document types first.
    # OPORD is last because it's frequently referenced inside other doc types
    # (e.g. "per OPORD 07-21") and would otherwise mis-classify AARs, SITREPs, etc.
    ("AAR",            r"\bAAR\b|\bAFTER[- ]ACTION\s+(?:REPORT|REVIEW)\b"),
    ("FRAGO",          r"\bFRAGO\b|\bFRAGMENTARY\s+ORDER\b"),
    ("WARNO",          r"\bWARNO\b|\bWARNING\s+ORDER\b"),
    ("SITREP",         r"\bSITREP\b|\bSITUATION\s+REPORT\b"),
    ("MEDEVAC",        r"\bMEDEVAC\s+REQUEST\b|\b9[- ]LINE\b"),
    ("MEDICAL_ANNEX",  r"ANNEX\s+\w+\s*\(.*MEDICAL|MEDICAL.*ANNEX|HEALTH\s+SERVICE\s+SUPPORT\s+ANNEX"),
    ("OPLAN",          r"\bOPLAN\b|\bOPERATION\s+PLAN\b"),
    ("OPORD",          r"\bOPORD\b|\bOPERATIONS\s+ORDER\b"),
]

def detect_doc_type(text):
    # Score each type by number of pattern hits in the first 5 000 chars
    sample = text[:5000]
    scores = {}
    for doc_type, pattern in DOC_TYPE_PATTERNS:
        hits = len(re.findall(pattern, sample, re.IGNORECASE))
        if hits:
            scores[doc_type] = hits
    if not scores:
        return "UNKNOWN"
    # Return the highest-priority type that has at least one hit
    for doc_type, _ in DOC_TYPE_PATTERNS:
        if doc_type in scores:
            return doc_type
    return "UNKNOWN"

# ── Format guides — teach the LLM what each document type contains ─────────────
FORMAT_GUIDES = {
    "OPORD": """\
OPORD (Operations Order) uses the standard 5-paragraph format:
1. SITUATION
   - classification, operation_name, effective_dtg, references, time_zone, issuing_unit
   - enemy_forces: composition, disposition, strength, capabilities, COAs
   - friendly_forces: higher/adjacent/supporting unit missions
   - attachments_detachments: units attached or detached with effective times
   - civil_considerations: ASCOPE factors affecting operations

2. MISSION
   - mission: verbatim mission statement (who/what/when/where/why — the 5 Ws)
   - commander_intent: purpose, key_tasks (list), end_state

3. EXECUTION
   - concept_of_operations: overall scheme of maneuver/fires
   - phases: list of phases with name, trigger, and description
   - tasks_to_subordinate_units: object per unit with its specific tasks
   - coordinating_instructions: instructions applying to 2+ units (PIR, movement, ROE, etc.)

4. SERVICE SUPPORT / SUSTAINMENT
   - sustainment_concept: logistics support concept
   - supply: supply class priorities, resupply procedures, Class VIII (medical)
   - transportation: available assets, priorities, routes
   - maintenance: maintenance priorities, deadlined equipment procedures
   - medical:
       receiving_hospitals: list with role level (Role 1/2/3/4) and location
       medevac_request_procedures: 9-line format and submission instructions
       medevac_platforms: aircraft/vehicle types and call signs
       casualty_collection_points: CCP locations and responsibilities
       blood_products: type O availability, walking blood bank procedures
       dental_support: dental unit/personnel assignments
       behavioral_health: mental health assets
       preventive_medicine: PVNTMED measures and contacts
       veterinary_support: food/water inspection procedures
       medical_logistics: MEDLOG, Class VIII resupply, push vs. pull
       mascal_trigger: threshold count and procedures
       medical_personnel: surgeons, medics, PA assignments
       succession_of_command: ordered list (medical chain)

5. COMMAND AND SIGNAL
   - succession_of_command: full ordered list
   - command_posts: CP locations (main, TAC, rear, jump)
   - signal: PACE plan (primary/alternate/contingency/emergency frequencies and call signs)
   - acknowledge: acknowledgement instructions""",

    "MEDICAL_ANNEX": """\
Medical Annex (commonly Annex H or similar) to an OPORD:
- annex_title: full annex title including parent OPORD
- issuing_unit: medical unit issuing the annex
- effective_dtg: effective date-time group
- references: maps, SOPs, other orders referenced
- situation:
    medical_threat: disease, non-battle injury risks, environmental hazards
    medical_enemy_situation: enemy medical capabilities/assets
    friendly_medical_assets: medical units/teams available
- mission: medical unit mission statement
- execution:
    medical_support_concept: overall approach
    receiving_hospitals: [{name, role_level, location, capacity, contact}]
    medevac_request_procedures: 9-line format, submission method, response time standards
    medevac_platforms: [{type, call_sign, location, availability}]
    casualty_collection_points: [{name, location_mgrs, unit_responsible}]
    treatment_facilities: aid stations, FSTs, CSHs with locations
    blood_products: type O whole blood, pRBCs, FFP procedures; walking blood bank SOP
    dental_support: dental unit and personnel
    behavioral_health: BH assets and referral procedures
    preventive_medicine: PVNTMED measures, contacts, environmental hazards
    veterinary_support: food/water inspection procedures
    medical_logistics: Class VIII push/pull, MEDLOG contact, resupply procedures
    medical_personnel: [{name, rank, role, unit}]
    mascal_trigger: threshold and activation procedures
    succession_of_command: [{rank_name, position}]
- sustainment: Class VIII supply, medical equipment maintenance
- command_and_signal:
    succession_of_command: ordered list
    medical_nets: frequencies and call signs
    command_posts: medical CP locations""",

    "AAR": """\
AAR (After Action Review/Report):
- unit: unit that conducted the event
- event_date: date of the event/operation
- event_location: where the event occurred
- event_type: type of operation or training event
- task_and_purpose: the assigned task and its purpose
- participants: units/personnel who participated
- timeline: sequence of key events with times
- sustains: list of things that went well (keep doing)
- improves: list of things that need improvement
- action_items: [{action, responsible_party, due_date, status}]
- lessons_learned: transferable lessons for future operations
- recommendations: recommendations for doctrine, training, or equipment
- casualties: any casualties incurred with circumstances
- equipment_issues: equipment problems encountered
- summary: overall event summary""",

    "FRAGO": """\
FRAGO (Fragmentary Order) — changes to an existing OPORD:
- references_parent_opord: the OPORD being modified
- effective_dtg: when changes take effect
- situation_changes: changes to situation paragraph
- mission_changes: changes to mission (if changed)
- execution_changes: changes to execution (tasks, phases, timing)
- sustainment_changes: logistics/medical changes
- command_signal_changes: comms or command changes
- coordinating_instructions: updated instructions""",

    "SITREP": """\
SITREP (Situation Report):
- reporting_unit: unit submitting
- dtg: date-time group of report
- location: unit location (MGRS)
- situation_summary: current tactical situation
- personnel_status: strength, casualties (KIA/WIA/MIA)
- equipment_status: operational readiness
- logistics_status: supply levels, resupply needs
- medical_status: sick/injured, medevac requests
- enemy_activity: observed enemy actions
- friendly_activity: own unit activities
- weather: current/forecast weather affecting ops
- next_report_dtg: when next report is due""",

    "UNKNOWN": """\
Extract all structured information from this military document.
Use the document's own section headers as JSON keys.
Capture every paragraph, list, and table as a field.""",
}

def tier2_ollama(full_text, fields, spacy_ents):
    already_found = {
        k: v["value"]
        for k, v in fields.items()
        if v.get("confidence", 0) >= TIER2_THRESHOLD and v["value"] is not None
        and v.get("method") not in ("spacy",)
    }

    doc_type = detect_doc_type(full_text)
    format_guide = FORMAT_GUIDES.get(doc_type, FORMAT_GUIDES["UNKNOWN"])
    print(f"[T2] detected doc type: {doc_type}")

    # llama3.1:8b supports 128k context — send as much text as safely fits
    text_limit = 28000
    doc_text = full_text[:text_limit]
    if len(full_text) > text_limit:
        # Append the tail too — end of documents often has command/signal sections
        doc_text += "\n\n[...]\n\n" + full_text[-4000:]

    prompt = f"""You are a military document analyst specialising in {doc_type} documents.

DOCUMENT FORMAT REFERENCE — {doc_type}:
{format_guide}

ALREADY CONFIRMED (include these values unchanged in your output):
{json.dumps(already_found, indent=2)}

YOUR TASK:
Extract EVERY piece of information from the document into a single flat JSON object.
- Use the field names from the format reference above where they apply.
- For any section or paragraph not covered by the reference, create a new snake_case key using the section header (e.g. "annex_a_fires", "tab_1_ccp_locations").
- Lists of items should be JSON arrays.
- Multi-paragraph text should be a single string.
- NEVER invent data. If a field is absent, omit it entirely (do not include null values).
- Return ONLY valid JSON. No markdown fences, no commentary.

DOCUMENT:
{doc_text}"""

    raw = ollama_chat(prompt, max_tokens=4096)

    schema = {}
    # Strategy 1: direct parse
    try:
        schema = json.loads(raw.strip())
    except Exception:
        pass

    # Strategy 2: extract outermost {...} block
    if not schema:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            try:
                schema = json.loads(m.group(0))
            except Exception:
                pass

    # Strategy 3: strip markdown fences then retry
    if not schema:
        cleaned = re.sub(r"```(?:json)?|```", "", raw).strip()
        m = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if m:
            try:
                schema = json.loads(m.group(0))
            except Exception:
                pass

    if not schema:
        print(f"[T2] JSON parse failed after all strategies, raw[:300]: {raw[:300]}")

    # Store detected doc type
    schema["doc_type"] = doc_type

    # Merge: regex high-confidence wins; Ollama fills everything else
    SPACY_OVERRIDEABLE = {"units", "personnel", "locations"}
    merged = {k: v for k, v in fields.items() if k not in SPACY_OVERRIDEABLE}

    for key, val in schema.items():
        if isinstance(val, str) and val in HALLUCINATED:
            continue
        if val is None:
            continue
        existing_conf = merged.get(key, {}).get("confidence", 0)
        existing_method = merged.get(key, {}).get("method", "")
        if existing_conf < TIER2_THRESHOLD or existing_method == "spacy":
            merged[key] = {"value": val, "confidence": 0.8, "method": "ollama"}

    # Re-add spaCy fields only if Ollama didn't cover them
    ollama_keys = {k for k, v in merged.items() if v.get("method") == "ollama"}
    if "locations" not in ollama_keys and "locations" in fields:
        merged["locations"] = fields["locations"]
    if "personnel" not in ollama_keys and "personnel" in fields:
        merged["personnel"] = fields["personnel"]

    return merged

# ── Storage helpers ────────────────────────────────────────────────────────────
def index_document(document_id, file_name, file_path, doc_type, fields, pages):
    # SQLite — document record
    with _db() as con:
        con.execute(
            "INSERT OR REPLACE INTO documents (id, file_name, file_path, doc_type, schema_json, extracted_at) "
            "VALUES (?,?,?,?,?,?)",
            (document_id, file_name, str(file_path), doc_type,
             json.dumps(fields), datetime.utcnow().isoformat()),
        )
        # SQLite — chunks
        for page in pages:
            chunk_id = str(uuid.uuid4())
            con.execute(
                "INSERT INTO chunks (id, document_id, page_num, text) VALUES (?,?,?,?)",
                (chunk_id, document_id, page["page_num"], page["text"]),
            )
        con.commit()

    # ChromaDB — embed chunks
    for page in pages:
        chunk_id = str(uuid.uuid4())
        try:
            chroma_col.add(
                documents=[page["text"]],
                ids=[chunk_id],
                metadatas=[{"document_id": document_id, "file_name": file_name, "page_num": page["page_num"]}],
            )
        except Exception as e:
            print(f"[ChromaDB] index error p{page['page_num']}: {e}")

# ── Main pipeline ──────────────────────────────────────────────────────────────
def run_pipeline(file_path, file_name, document_id=None):
    pages = read_document(file_path)
    if not pages:
        return {"fields": {}, "pages": 0, "text_length": 0}

    full_text = "\n\n".join(p["text"] for p in pages)

    # Tier 0 — regex
    fields = tier0_regex(full_text)
    t0 = sum(1 for v in fields.values() if v["value"] is not None)
    print(f"[T0] {t0}/{len(fields)} fields from regex")

    # Tier 1 — spaCy
    fields, spacy_ents = tier1_spacy(full_text, fields)
    t1 = sum(1 for v in fields.values() if v.get("method") == "spacy")
    print(f"[T1] spaCy added {t1} field(s)")

    # Tier 2 — Ollama (always runs — it proposes the adaptive schema)
    fields = tier2_ollama(full_text, fields, spacy_ents)
    t2 = sum(1 for v in fields.values() if v.get("method") == "ollama")
    print(f"[T2] Ollama filled {t2} field(s) | total schema fields: {len(fields)}")

    # Persist to SQLite + ChromaDB
    print(f"[index] document_id={document_id}, pages={len(pages)}")
    if document_id:
        doc_type = fields.get("doc_type", {}).get("value") if isinstance(fields.get("doc_type"), dict) else fields.get("doc_type") or "unknown"
        index_document(document_id, file_name, file_path, str(doc_type), fields, pages)
        print(f"[index] done — {len(pages)} chunks added to ChromaDB")
    else:
        print("[index] SKIPPED — no document_id received")

    return {"fields": fields, "pages": len(pages), "text_length": len(full_text)}

# ── Flask routes ───────────────────────────────────────────────────────────────

@app.route("/extract-pdf", methods=["POST"])
def extract_pdf():
    try:
        data = request.json
        file_path = data.get("file_path", "")
        document_id = data.get("document_id") or None
        file_name = data.get("file_name", Path(file_path).name)
        print(f"[extract-pdf] file={file_name} document_id={document_id!r} keys={list(data.keys())}")

        if not file_path or not Path(file_path).exists():
            return jsonify({"error": f"File not found: {file_path}"}), 404

        result = run_pipeline(file_path, file_name, document_id)
        summary = {
            k: {"value": v["value"], "confidence": v["confidence"], "method": v["method"]}
            for k, v in result["fields"].items()
        }
        return jsonify({**result, "summary": summary})

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/extract", methods=["POST"])
def extract_text():
    """Text-only extraction (called by Node for .txt/.md files)."""
    try:
        data = request.json
        text = data.get("text", "")
        if not text:
            return jsonify({"error": "text is required"}), 400

        fields = tier0_regex(text)
        fields, spacy_ents = tier1_spacy(text, fields)
        fields = tier2_ollama(text, fields, spacy_ents)
        return jsonify({k: v["value"] for k, v in fields.items()})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _get_schema_for_doc(document_id):
    """Load the saved schema JSON from SQLite for a document."""
    with _db() as con:
        row = con.execute(
            "SELECT schema_json, file_name FROM documents WHERE id=?", (document_id,)
        ).fetchone()
    if not row or not row[0]:
        return None, None
    try:
        return json.loads(row[0]), row[1]
    except Exception:
        return None, row[1]

def _schema_to_text(schema, file_name, doc_id=None):
    """Convert a schema dict to readable key: value lines for the LLM.

    doc_id is passed explicitly from the caller (the SQLite primary key) so
    the function works correctly even when the stored schema_json has no _meta.
    """
    # Prefer the caller-supplied doc_id; fall back to _meta if present
    if not doc_id:
        doc_id = schema.get("_meta", {}).get("documentId") or "UNKNOWN"
    lines = [f"Document: {file_name} [ID: {doc_id}]"]
    fields = schema.get("fields", schema)
    for key, field in fields.items():
        if key == "_meta":
            continue
        if isinstance(field, dict):
            val = field.get("value")
        else:
            val = field
        if val is None:
            continue
        if isinstance(val, (dict, list)):
            val = json.dumps(val, ensure_ascii=False)
        lines.append(f"  {key}: {val}")
    return "\n".join(lines)

META_QUERY_PATTERNS = re.compile(
    r"\b(what (data|documents?|files?|reports?|info) (do you have|are (in|available)|have you (got|indexed))|"
    r"(list|show|summarize|give me).{0,20}(documents?|files?|reports?|data)|"
    r"what('s| is) in (the )?(system|database|knowledge base)|"
    r"do you have (any )?(data|documents?|files?|reports?)|"
    r"what (have you|do you) (know|have)|"
    r"(available|indexed) documents?)\b",
    re.IGNORECASE,
)

def _build_inventory_response():
    """Return a human-readable summary of every document currently in SQLite."""
    with _db() as con:
        rows = con.execute(
            "SELECT id, file_name, doc_type, schema_json, extracted_at FROM documents ORDER BY extracted_at DESC"
        ).fetchall()

    if not rows:
        return "No documents are currently indexed in the knowledge base.", []

    lines = [f"I have **{len(rows)} document(s)** indexed:\n"]
    citations = []
    for doc_id, file_name, doc_type, schema_json, extracted_at in rows:
        schema = {}
        try:
            schema = json.loads(schema_json or "{}")
        except Exception:
            pass

        fields = schema.get("fields", schema)
        def fv(key):
            f = fields.get(key, {})
            return f.get("value") if isinstance(f, dict) else f

        title    = fv("title") or file_name
        date_val = fv("date")
        date_str = date_val[:10] if date_val and len(date_val) >= 10 else (date_val or "unknown date")
        location = fv("location") or fv("mgrsGrid") or "unknown location"
        dtype    = fv("doc_type") or doc_type or "unknown type"
        mission  = fv("mission")

        summary = f"- **{doc_id}** — {file_name}\n"
        summary += f"  Type: {dtype} | Date: {date_str} | Location: {location}\n"
        if title and title != file_name:
            summary += f"  Title: {title}\n"
        if mission:
            snippet = mission[:120].replace("\n", " ")
            summary += f"  Mission: {snippet}{'...' if len(mission) > 120 else ''}\n"

        lines.append(summary)
        citations.append({"document_id": doc_id, "file_name": file_name})

    return "\n".join(lines), citations


@app.route("/chat", methods=["POST"])
def chat():
    try:
        data = request.json
        user_message = data.get("message", "")
        if not user_message:
            return jsonify({"error": "Message is required"}), 400

        # ── Meta-query: user asking what's in the system ──────────────────────
        if META_QUERY_PATTERNS.search(user_message):
            response_text, citations = _build_inventory_response()
            return jsonify({"response": response_text, "citations": citations})

        # ── Normal RAG query ──────────────────────────────────────────────────
        try:
            doc_count = chroma_col.count()
        except Exception:
            doc_count = 0

        if doc_count == 0:
            return jsonify({"response": "No documents are currently indexed. Please upload a report first.", "citations": []}), 200

        n_results = min(5, doc_count)
        try:
            results = chroma_col.query(query_texts=[user_message], n_results=n_results)
        except Exception as qe:
            print(f"[chat] ChromaDB query error: {qe}")
            return jsonify({"response": "Search index unavailable. Please try again.", "citations": []}), 200

        if not results["documents"] or not results["documents"][0]:
            return jsonify({"response": "No relevant documents found in the knowledge base.", "citations": []}), 200

        metas = results["metadatas"][0]

        # Collect unique document IDs from search results
        seen_docs = {}
        for meta in metas:
            doc_id = meta.get("document_id")
            if doc_id and doc_id not in seen_docs:
                seen_docs[doc_id] = meta.get("file_name", "Unknown")

        # Build context from schema JSON — reflects user edits via JSON panel
        context_parts, citations = [], {}
        for doc_id, file_name in seen_docs.items():
            schema, fname = _get_schema_for_doc(doc_id)
            display_name = fname or file_name
            if schema:
                context_parts.append(_schema_to_text(schema, display_name, doc_id))
                citations[doc_id] = {"document_id": doc_id, "file_name": display_name}
            else:
                for chunk, meta in zip(results["documents"][0], metas):
                    if meta.get("document_id") == doc_id:
                        context_parts.append(f"[{display_name} | ID: {doc_id}]: {chunk}")
                citations[doc_id] = {"document_id": doc_id, "file_name": display_name}

        if not context_parts:
            return jsonify({"response": "No relevant documents found in the knowledge base.", "citations": []}), 200

        real_ids = ", ".join(seen_docs.keys())
        prompt = (
            f"The only valid document IDs in the knowledge base are: {real_ids}\n"
            "Answer the question using ONLY the structured document data below.\n"
            "After every fact cite the EXACT document ID in brackets, e.g. [RPT-XXXXXX]. "
            "If the fact comes from a named field also include it, e.g. [RPT-XXXXXX · mission]. "
            "Never invent or guess a document ID.\n\n"
            f"DOCUMENT DATA:\n{chr(10).join(context_parts)}\n\n"
            f"Question: {user_message}"
        )
        answer = ollama_chat(
            prompt,
            system=(
                "You are a military medical operations assistant. "
                "Answer using only the provided document data. "
                "Always cite facts with the real document ID given in the DOCUMENT DATA header — "
                "never use placeholder IDs like RPT-ABC123 or RPT-XXXXXX. "
                "Be concise and direct."
            ),
            max_tokens=1024,
        )
        return jsonify({"response": answer, "citations": list(citations.values())})

    except Exception as e:
        print(f"[chat] error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/documents", methods=["GET"])
def list_documents():
    with _db() as con:
        rows = con.execute("SELECT id, file_name, doc_type, extracted_at FROM documents").fetchall()
    return jsonify([{"id": r[0], "file_name": r[1], "doc_type": r[2], "extracted_at": r[3]} for r in rows])


@app.route("/reindex", methods=["POST"])
def reindex():
    """
    Re-index a document into ChromaDB after its extraction JSON has been edited.
    Body: { "document_id": "...", "file_path": "...", "file_name": "...", "fields": {...} }
    """
    try:
        data = request.json
        document_id = data.get("document_id")
        file_path   = data.get("file_path")
        file_name   = data.get("file_name", "")
        fields      = data.get("fields", {})

        if not document_id or not file_path:
            return jsonify({"error": "document_id and file_path are required"}), 400

        # Re-read pages from the original file
        pages = read_document(file_path)
        if not pages:
            return jsonify({"error": "Could not read file"}), 400

        # Pull updated file_name from _meta.fileName if user edited it
        updated_file_name = file_name
        if isinstance(fields, dict):
            meta_name = fields.get("_meta", {}).get("fileName")
            if meta_name:
                updated_file_name = meta_name

        # Delete existing ChromaDB chunks for this document
        existing = chroma_col.get(where={"document_id": document_id})
        if existing and existing["ids"]:
            chroma_col.delete(ids=existing["ids"])
            print(f"[reindex] deleted {len(existing['ids'])} old chunks for {document_id}")

        # Re-add chunks with updated metadata (use corrected file_name)
        for page in pages:
            chunk_id = str(uuid.uuid4())
            chroma_col.add(
                documents=[page["text"]],
                ids=[chunk_id],
                metadatas=[{"document_id": document_id, "file_name": updated_file_name, "page_num": page["page_num"]}],
            )

        # Update SQLite schema and file_name with corrected values
        with _db() as con:
            con.execute(
                "UPDATE documents SET schema_json=?, file_name=?, extracted_at=? WHERE id=?",
                (json.dumps(fields), updated_file_name, datetime.utcnow().isoformat(), document_id),
            )
            con.commit()
        print(f"[reindex] file_name updated to: {updated_file_name}")

        print(f"[reindex] {document_id} re-indexed with {len(pages)} chunks")
        return jsonify({"message": f"Re-indexed {len(pages)} chunks"})

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/delete-document", methods=["POST"])
def delete_document():
    """Remove all ChromaDB chunks and the SQLite record for a document."""
    try:
        data = request.json
        document_id = data.get("document_id")
        if not document_id:
            return jsonify({"error": "document_id is required"}), 400

        # Remove ChromaDB chunks
        existing = chroma_col.get(where={"document_id": document_id})
        deleted_chunks = 0
        if existing and existing["ids"]:
            chroma_col.delete(ids=existing["ids"])
            deleted_chunks = len(existing["ids"])
            print(f"[delete] removed {deleted_chunks} ChromaDB chunks for {document_id}")

        # Remove SQLite record and chunks
        with _db() as con:
            con.execute("DELETE FROM chunks WHERE document_id=?", (document_id,))
            con.execute("DELETE FROM documents WHERE id=?", (document_id,))
            con.commit()
        print(f"[delete] removed SQLite record for {document_id}")

        return jsonify({"message": f"Deleted {deleted_chunks} chunks", "document_id": document_id})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/health", methods=["GET"])
def health():
    try:   chunk_count = chroma_col.count()
    except: chunk_count = 0
    with _db() as con:
        doc_count = con.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
    return jsonify({"status": "ok", "documents": doc_count, "chunks_indexed": chunk_count})


if __name__ == "__main__":
    print("✅ Pipeline ready — Tier 0 (regex) | Tier 1 (spaCy) | Tier 2 (Ollama llama3.1:8b)")
    app.run(host="0.0.0.0", port=5002, debug=True)
